# -*- coding: utf-8 -*-
import os
import re
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Reconfigure stdout for UTF-8 encoding to avoid Windows console errors
if sys.stdout.encoding != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except AttributeError:
        pass


# Paths
DOC_DIR = r"c:\Users\Kanomjean\Downloads\project\KnowledgeBase\07_document"
FLOWCHARTS_DIR = r"c:\Users\Kanomjean\Downloads\project\flowcharts"
OUTPUT_PATH = os.path.join(DOC_DIR, "Chapter3_Methodology.docx")

def set_font(run, font_name='TH Sarabun New', size_pt=16, bold=False, italic=False, color=None):
    run.font.name = font_name
    rFonts = run._r.get_or_add_rPr().get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def add_markdown_paragraph(doc, text, style=None, align=None, indent_first=False, space_before=0, space_after=6):
    p = doc.add_paragraph()
    if style:
        p.style = style
    if align:
        p.alignment = align
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.25)
    
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            set_font(r, bold=True)
        elif part.startswith('*') and part.endswith('*'):
            r = p.add_run(part[1:-1])
            set_font(r, italic=True)
        else:
            if part:
                r = p.add_run(part)
                set_font(r)
    return p

def main():
    print("📖 Starting conversion of Chapter 3 markdown report to formatted Word document...")
    doc = Document()
    
    # Configure margins (RMUTSB standard)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        
        # Configure footer for page numbers
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer_p.add_run()
        set_font(footer_run, size_pt=12)
        add_page_number(footer_run)

    # Define Chapter 3 file path
    chapter3_path = os.path.join(DOC_DIR, "Chapter3_Methodology.md")
    
    if not os.path.exists(chapter3_path):
        print(f"❌ File not found: {chapter3_path}")
        sys.exit(1)
        
    print(f"🔄 Processing Chapter 3...")
    with open(chapter3_path, "r", encoding="utf-8") as f:
        content = f.read()
        
    # Clean guidelines/header notes from chapter
    content = re.sub(r"^# คู่มือการจัดรูปแบบหน้ากระดาษ.*?\n---\n", "", content, flags=re.DOTALL | re.IGNORECASE)
    
    # Split into blocks (paragraphs, tables, lists, code blocks)
    lines = content.split("\n")
    
    in_table = False
    table_lines = []
    in_code_block = False
    code_lines = []
    first_chapter = True
    
    for line in lines:
        striped_line = line.strip()
        
        # Code block toggle
        if striped_line.startswith("```"):
            if in_code_block:
                in_code_block = False
                # Write code lines as a styled paragraph
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.right_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(code_text)
                set_font(run, font_name="Courier New", size_pt=10, color=RGBColor(80, 80, 80))
                
                code_lines = []
            else:
                in_code_block = True
            continue
            
        if in_code_block:
            code_lines.append(line)
            continue
            
        # Table parse toggle
        if striped_line.startswith("|"):
            in_table = True
            table_lines.append(striped_line)
            continue
        else:
            if in_table:
                in_table = False
                # Render table
                render_word_table(doc, table_lines)
                table_lines = []
        
        # Skip empty lines
        if not striped_line:
            continue
            
        # Process Headings
        if striped_line.startswith("# ") or (striped_line.startswith("## ") and "บทที่" in striped_line):
            heading_text = striped_line.lstrip("#").strip()
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(36)  # Simulate 1.5 inch top margin
            p.paragraph_format.space_after = Pt(18)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(heading_text)
            set_font(run, size_pt=20, bold=True)
            
        elif striped_line.startswith("## "):
            # Heading 1 (18pt Bold)
            heading_text = striped_line.lstrip("#").strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(heading_text)
            set_font(run, size_pt=18, bold=True)
            
        elif striped_line.startswith("### "):
            # Heading 2 (18pt Bold)
            heading_text = striped_line.lstrip("#").strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(heading_text)
            set_font(run, size_pt=18, bold=True)
            
        elif striped_line.startswith("#### "):
            # Heading 3 (16pt Bold)
            heading_text = striped_line.lstrip("#").strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(heading_text)
            set_font(run, size_pt=16, bold=True)
            
        elif striped_line.startswith("##### "):
            # Heading 4 (16pt Bold/Italic)
            heading_text = striped_line.lstrip("#").strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(heading_text)
            set_font(run, size_pt=16, bold=True, italic=True)
            
        # Lists
        elif striped_line.startswith("* ") or striped_line.startswith("- "):
            list_text = striped_line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            
            # Format runs
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', list_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    set_font(r, bold=True)
                elif part.startswith('*') and part.endswith('*'):
                    r = p.add_run(part[1:-1])
                    set_font(r, italic=True)
                else:
                    if part:
                        r = p.add_run(part)
                        set_font(r)
                        
        elif re.match(r'^\d+\)\s|^\d+\.\s', striped_line):
            # Numbered list items
            list_text = re.sub(r'^\d+[\)\.]\s', '', striped_line).strip()
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            
            # Format runs
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', list_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    set_font(r, bold=True)
                elif part.startswith('*') and part.endswith('*'):
                    r = p.add_run(part[1:-1])
                    set_font(r, italic=True)
                else:
                    if part:
                        r = p.add_run(part)
                        set_font(r)
        
        # Check for image references and placeholders
        elif "ภาพที่" in striped_line and ("**ภาพที่" in striped_line or "ภาพที่ 3." in striped_line):
            image_injected = False
            
            # Identify diagram to inject
            if "ภาพที่ 3.1" in striped_line:
                image_injected = inject_image(doc, "01_system_overview.png")
            elif "ภาพที่ 3.2" in striped_line:
                image_injected = inject_image(doc, "02_auth_and_profile_flow.png")
            elif "ภาพที่ 3.3" in striped_line:
                image_injected = inject_image(doc, "04_hardware_compatibility_check_flow.png")
            elif "ภาพที่ 3.4" in striped_line:
                image_injected = inject_image(doc, "03_product_listing_and_antifraud_flow.png")
            elif "ภาพที่ 3.5" in striped_line:
                image_injected = inject_image(doc, "06_c2c_order_booking_and_chat_flow.png")
            
            # If image is injected, let's write the caption centered
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(12)
            
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', striped_line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    set_font(r, bold=True)
                elif part.startswith('*') and part.endswith('*'):
                    r = p.add_run(part[1:-1])
                    set_font(r, italic=True)
                else:
                    if part:
                        r = p.add_run(part)
                        set_font(r)
                        
        # Skip horizontal rules
        elif striped_line == "---":
            continue
            
        # Skip placeholder brackets
        elif striped_line.startswith("*(") and striped_line.endswith(")*"):
            continue
        elif striped_line.startswith("(") and striped_line.endswith(")"):
            continue
            
        # Normal paragraph
        else:
            add_markdown_paragraph(doc, striped_line, align=WD_ALIGN_PARAGRAPH.LEFT, indent_first=True)

    # Save document
    print(f"💾 Saving Word document to {OUTPUT_PATH}...")
    doc.save(OUTPUT_PATH)
    print("✅ Conversion completed successfully!")

def inject_image(doc, filename):
    img_path = os.path.join(FLOWCHARTS_DIR, filename)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(5.5))
        return True
    else:
        print(f"⚠️ Image not found: {img_path}")
        return False

def render_word_table(doc, table_lines):
    rows_data = []
    for line in table_lines:
        cells = [c.strip() for c in line.split("|")]
        if cells and cells[0] == "":
            cells.pop(0)
        if cells and cells[-1] == "":
            cells.pop()
        
        if cells and all(re.match(r'^[\s:-]+$', c) for c in cells):
            continue
            
        if cells:
            rows_data.append(cells)
            
    if not rows_data:
        return
        
    num_cols = max(len(r) for r in rows_data)
    num_rows = len(rows_data)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    border_style = {"sz": 4, "val": "single", "color": "BFBFBF", "space": "0"}
    
    for r_idx, row_cells in enumerate(rows_data):
        row = table.rows[r_idx]
        is_header = (r_idx == 0)
        
        trPr = row._tr.get_or_add_trPr()
        if is_header:
            trPr.append(OxmlElement('w:tblHeader'))
            
        for c_idx, cell_value in enumerate(row_cells):
            if c_idx < len(row.cells):
                cell = row.cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.0
                
                if is_header or c_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', cell_value)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        r = p.add_run(part[2:-2])
                        set_font(r, size_pt=14, bold=True)
                    elif part.startswith('*') and part.endswith('*'):
                        r = p.add_run(part[1:-1])
                        set_font(r, size_pt=14, italic=True)
                    else:
                        if part:
                            r = p.add_run(part)
                            set_font(r, size_pt=14, bold=is_header)
                
                set_cell_border(cell, top=border_style, bottom=border_style, left=border_style, right=border_style)
                if is_header:
                    set_cell_background(cell, "F2F2F2")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

if __name__ == "__main__":
    main()
